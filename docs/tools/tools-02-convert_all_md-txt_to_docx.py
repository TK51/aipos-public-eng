# convert_all_md-txt_to_docx.py

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
import re

script_dir = Path(__file__).parent
text_files = list(script_dir.glob("*.md")) + list(script_dir.glob("*.txt"))

if not text_files:
    print("⚠️ No .md or .txt files found in folder.")
    exit()

for text_file in text_files:
    content = text_file.read_text(encoding="utf-8")
    docx_file = text_file.with_suffix(".docx")
    doc = Document()

    # Set narrow margins: 0.5 inch all around
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Force 'Normal' style to compact spacing and Calibri
    normal_style = doc.styles['Normal']
    normal_format = normal_style.paragraph_format
    normal_format.space_after = Pt(0)
    normal_format.space_before = Pt(0)
    normal_format.line_spacing = 1
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # Compact spacing for headings (Heading 1–3)
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        para_format = heading_style.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)
        para_format.line_spacing = 1
        heading_style.font.name = "Calibri"
        heading_style.font.size = Pt(12 + (3 - level))  # H1 = 14, H2 = 13, H3 = 12

    # Define base custom style
    styles = doc.styles
    if "MarkdownBase" not in styles:
        style = styles.add_style("MarkdownBase", WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    # Title (file name as Heading 0)
    doc.add_heading(text_file.stem.replace("_", " "), 0)

    for line in content.splitlines():
        stripped = line.strip()

        # Empty line
        if not stripped:
            para = doc.add_paragraph(style="MarkdownBase")
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        # Bullets (both - and •)
        if line.lstrip().startswith("•") or line.lstrip().startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1
            para.style.font.name = "Calibri"

            bullet_text = line.lstrip()[1:].strip() if line.lstrip().startswith("•") else line.lstrip()[2:].strip()

            parts = re.split(r"(\*\*.*?\*\*)", bullet_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
            continue

        # Regular paragraph with optional bold
        para = doc.add_paragraph(style="MarkdownBase")
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1
        para.style.font.name = "Calibri"

        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    doc.save(docx_file)
    print(f"✅ Converted: {text_file.name} → {docx_file.name}")



#python #markdownconverter #docxgenerator #textprocessing #automation #reportingautomation #workflowtools #documentationtools #traceability #businessdataanalyst

#aiposbuilt #fromukrainianswithlovetohumankind 🇺🇦
